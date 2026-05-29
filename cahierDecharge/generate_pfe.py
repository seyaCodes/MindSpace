#!/usr/bin/env python3
"""
Generate complete MindSpace PFE report as a .docx file.
University Abdelhamid Ibn Badis – Mostaganem, 2025–2026
"""

import os
import copy
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR  = r'd:\program\mind_space\cahierDecharge'
IMG_DIR   = os.path.join(BASE_DIR, 'extracted_images')
DIAG_DIR  = os.path.join(BASE_DIR, 'generated_diagrams')
TMPL_PATH = os.path.join(BASE_DIR, 'Template_Rapport_Licence_2025_2026.docx')
OUT_PATH  = os.path.join(BASE_DIR, 'MindSpace_PFE_Report.docx')

# ── helpers ──────────────────────────────────────────────────

def new_doc():
    doc = Document(TMPL_PATH)
    # Clear body content while preserving the sectPr (section properties) element
    body = doc.element.body
    # Save the sectPr element before clearing
    sectPr = body.find(qn('w:sectPr'))
    # Remove all children
    for el in list(body):
        body.remove(el)
    # Re-append sectPr so sections work correctly
    if sectPr is not None:
        body.append(sectPr)
    else:
        # Create a minimal sectPr
        sectPr = OxmlElement('w:sectPr')
        body.append(sectPr)
    # Set margins per template spec
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.0)
    return doc


def set_para_fmt(para, font_name='Times New Roman', font_size=12,
                 bold=False, italic=False, color=None,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                 space_after=Pt(6), line_spacing=1.5,
                 first_line_indent=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_after = space_after
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    if first_line_indent is not None:
        pf.first_line_indent = first_line_indent
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor(*color)


def body(doc, text, bold=False, italic=False, indent=True):
    p = doc.add_paragraph(style='Texte mémoire')
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = bold
    run.font.italic = italic
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.5
    if indent:
        pf.first_line_indent = Cm(1.25)
    return p


def heading1(doc, text):
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p


def heading2(doc, text):
    p = doc.add_heading(text, level=2)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p


def heading3(doc, text):
    p = doc.add_heading(text, level=3)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p


def heading4(doc, text):
    p = doc.add_heading(text, level=4)
    for run in p.runs:
        run.font.name = 'Times New Roman'
    return p


def add_figure(doc, img_path, caption, width_cm=14.0):
    if img_path and os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        p = doc.add_paragraph(f'[Figure: {caption}]')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in cap.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.font.bold = True
    cap.paragraph_format.space_after = Pt(12)
    return p


def add_page_break(doc):
    doc.add_page_break()


def centered(doc, text, size=12, bold=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Paragraph')
    run = p.add_run(('• ' if level == 0 else '  – ') + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    return p


def requirement_table(doc, rows):
    """rows = list of (id, description, priority)"""
    table = doc.add_table(rows=len(rows)+1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ['ID', 'Description', 'Priority']):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, (rid, desc, prio) in enumerate(rows):
        row = table.rows[i+1].cells
        for cell, txt in zip(row, [rid, desc, prio]):
            cell.text = txt
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    return table


def uc_table(doc, fields):
    """fields = list of (label, value)"""
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(fields):
        r = table.rows[i].cells
        r[0].text = label
        r[0].paragraphs[0].runs[0].font.bold = True
        r[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        r[0].paragraphs[0].runs[0].font.size = Pt(10)
        r[0].width = Cm(4)
        r[1].text = value
        r[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        r[1].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    return table


def sd_table(doc, fields):
    """Sequence diagram textual description table."""
    return uc_table(doc, fields)


def abbr_table(doc, rows):
    table = doc.add_table(rows=len(rows)+1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, ['Abbreviation', 'Full Expression']):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (abbr, full) in enumerate(rows):
        row = table.rows[i+1].cells
        row[0].text = abbr
        row[0].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row[0].paragraphs[0].runs[0].font.size = Pt(10)
        row[1].text = full
        row[1].paragraphs[0].runs[0].font.name = 'Times New Roman'
        row[1].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()
    return table


def comparison_table(doc):
    headers = ['Criteria', 'Wysa', 'Woebot', 'Daylio', 'Mind Space']
    rows = [
        ('Conversational AI', '✓', '✓', '✗', '✓'),
        ('EFT-informed responses', '✗', '✗', '✗', '✓'),
        ('Automatic theme clustering', '✗', '✗', '✗', '✓'),
        ('Longitudinal arc analysis', '✗', '✗', 'Partial', '✓'),
        ('Safety classifier (pre-gate)', '✗', 'Partial', '✗', '✓'),
        ('Crisis resources (localised)', 'Partial', 'Partial', '✗', '✓ (Algeria)'),
        ('No advertising data sharing', 'Partial', '✗', 'Partial', '✓'),
        ('Streaming token-by-token', '✗', '✗', 'N/A', '✓'),
        ('Open-source / free tier', '✗', '✗', 'Freemium', '✓ (v1)'),
    ]
    table = doc.add_table(rows=len(rows)+1, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for cell, txt in zip(hdr, headers):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for i, row_data in enumerate(rows):
        row = table.rows[i+1].cells
        for cell, txt in zip(row, row_data):
            cell.text = txt
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


def db_schema_table(doc, table_name, columns):
    """columns = list of (name, type, constraints, description)"""
    t = doc.add_table(rows=len(columns)+1, cols=4)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for cell, txt in zip(hdr, ['Column', 'Type', 'Constraints', 'Notes']):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for i, (col, typ, cons, note) in enumerate(columns):
        row = t.rows[i+1].cells
        for cell, txt in zip(row, [col, typ, cons, note]):
            cell.text = txt
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()
    return t


# ── image paths ──────────────────────────────────────────────
def img(name):
    p = os.path.join(IMG_DIR, name)
    return p if os.path.exists(p) else None

def diag(name):
    p = os.path.join(DIAG_DIR, name)
    return p if os.path.exists(p) else None


# ════════════════════════════════════════════════════════════════
# BUILD DOCUMENT
# ════════════════════════════════════════════════════════════════
def build():
    doc = new_doc()

    # ── TITLE PAGE ────────────────────────────────────────────
    centered(doc, 'Ministry of Higher Education and Scientific Research', 11)
    centered(doc, 'University Abdelhamid Ibn Badis – Mostaganem', 11)
    centered(doc, 'Faculty of Exact Sciences and Computer Science', 11)
    centered(doc, 'Department of Mathematics and Computer Science', 11)
    centered(doc, 'Program: Computer Science', 11)
    doc.add_paragraph()
    centered(doc, 'License Project Report in Computer Science', 14, bold=True)
    centered(doc, 'Speciality: Computer Systems', 12, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, 'Title:', 12)
    centered(doc, 'Mind Space', 20, bold=True)
    centered(doc, 'An AI-Powered Mobile Application for Guided Emotional Self-Reflection', 14, bold=True)
    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('Student:\t\tSeya [Family Name]\n')
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    run2 = p.add_run('Supervisor:\t\tDr. Bentaouza')
    run2.font.name = 'Times New Roman'; run2.font.size = Pt(12)

    doc.add_paragraph()
    doc.add_paragraph()
    centered(doc, 'Academic Year 2025–2026', 12, bold=True)

    add_page_break(doc)

    # ── ACKNOWLEDGEMENTS ──────────────────────────────────────
    centered(doc, 'Acknowledgements', 16, bold=True)
    doc.add_paragraph()
    body(doc,
         'We would like to express our sincere gratitude to our supervisor, Dr. Bentaouza, '
         'for his guidance, availability, and constructive feedback throughout this project. '
         'His expertise and encouragement were invaluable at every stage of development.')
    body(doc,
         'We also extend our thanks to the faculty of the Department of Mathematics and '
         'Computer Science at the University of Mostaganem for their teaching and support '
         'throughout the licence programme.')
    body(doc,
         'Finally, we thank all those who, directly or indirectly, contributed to the '
         'realisation of this work.')
    add_page_break(doc)

    # ── DEDICATION ────────────────────────────────────────────
    centered(doc, 'Dedication', 16, bold=True)
    doc.add_paragraph()
    body(doc, 'To my family, whose unwavering support and encouragement made this journey possible.')
    body(doc, 'To everyone who has ever felt unseen in their emotions — this work is for you.')
    add_page_break(doc)

    # ── ABSTRACT (EN) ────────────────────────────────────────
    centered(doc, 'Abstract', 14, bold=True)
    doc.add_paragraph()
    body(doc,
         'Mind Space is a mobile application for AI-assisted emotional self-reflection, designed '
         'to help users explore their inner world through guided conversation with an AI companion '
         'named Sage. The application implements an Emotion-Focused Therapy (EFT)-informed '
         'conversational framework, a safety-first crisis detection pipeline, and an automatic '
         'thematic clustering mechanism called the Arc system, which groups recurring emotional '
         'themes across sessions using semantic embedding and cosine similarity. The system is '
         'built as an Android application using Flutter and Riverpod, supported by a serverless '
         'backend on Supabase Edge Functions, PostgreSQL with pgvector for vector similarity '
         'search, and three AI service providers: Groq (LLaMA 3.3 70B for conversation, '
         'LLaMA 3.1 8B for safety classification), OpenAI (text-embedding-3-small), and '
         'OpenRouter (DeepSeek R1 for longitudinal arc insight generation). This work addresses '
         'the challenge of making emotional self-reflection accessible, structured, and private '
         'for individuals who face barriers to formal mental health support. The main '
         'contributions are: a complete implementation of a mobile mental health companion '
         'application, an EFT-informed prompt engineering methodology, and an embedding-based '
         'automatic emotional theme detection system with a calibrated similarity threshold of 0.78.')
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Keywords: ')
    run.font.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    run2 = p.add_run(
        'mobile application, emotional self-reflection, artificial intelligence, '
        'large language models, Emotion-Focused Therapy, semantic embeddings, '
        'cosine similarity, Flutter, Supabase, pgvector.')
    run2.font.name = 'Times New Roman'; run2.font.size = Pt(12)
    add_page_break(doc)

    # ── ABSTRACT (FR) ────────────────────────────────────────
    centered(doc, 'Résumé', 14, bold=True)
    doc.add_paragraph()
    body(doc,
         "Mind Space est une application mobile d'accompagnement emotionnel guide par intelligence "
         "artificielle, concue pour aider les utilisateurs a explorer leur vie interieure a travers "
         "des conversations guidees avec un compagnon IA nomme Sage. L'application met en oeuvre un "
         "cadre conversationnel inspire de la Therapie Centree sur les Emotions (TCE), un pipeline "
         "de detection de crise axe sur la securite, et un mecanisme de regroupement thematique "
         "automatique appele le systeme d'Arcs, qui regroupe les themes emotionnels recurrents des "
         "sessions a l'aide d'encodages semantiques et de la similarite cosinus. Le systeme est "
         "developpe en tant qu'application Android avec Flutter et Riverpod, soutenu par un backend "
         "entierement sans serveur base sur Supabase Edge Functions, PostgreSQL avec pgvector, et "
         "trois fournisseurs de services d'IA incluant Groq (LLaMA 3.3 70B), OpenAI "
         "(text-embedding-3-small) et OpenRouter (DeepSeek R1). Ce travail aborde le defi de rendre "
         "l'auto-reflexion emotionnelle accessible, structuree et privee pour les individus qui font "
         "face a des obstacles dans l'acces au soutien en sante mentale formelle.")
    doc.add_paragraph()
    p = doc.add_paragraph()
    run = p.add_run('Mots-clés : ')
    run.font.bold = True; run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    run2 = p.add_run(
        'application mobile, auto-réflexion émotionnelle, intelligence artificielle, '
        'grands modèles de langage, Thérapie Centrée sur les Émotions, encodages sémantiques, '
        'Flutter, Supabase, pgvector.')
    run2.font.name = 'Times New Roman'; run2.font.size = Pt(12)
    add_page_break(doc)

    # ── ABSTRACT (AR) ────────────────────────────────────────
    centered(doc, 'ملخص', 14, bold=True)
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(
        'ماينـد سبيس هو تطبيق جوّال للتأمل العاطفي الذاتي المدعوم بالذكاء الاصطناعي، '
        'صُمِّم لمساعدة المستخدمين على استكشاف عوالمهم الداخلية من خلال محادثات موجَّهة مع مرافق '
        'ذكاء اصطناعي يُدعى "سيج". يعتمد التطبيق إطارًا حواريًا مستوحى من العلاج المرتكز على '
        'المشاعر (EFT)، ونظام كشف مسبق للأزمات، وآلية تجميع تلقائية للمحاور العاطفية تُسمى نظام '
        '"الأقواس"، التي تُجمِّع الموضوعات العاطفية المتكررة عبر الجلسات باستخدام التمثيلات '
        'الدلالية وقياس التشابه الجيبي. يعمل النظام كتطبيق Android مبني بإطار Flutter ومنصة '
        'Supabase اللامركزية، مع ثلاثة مزوِّدي خدمات ذكاء اصطناعي: Groq وOpenAI وOpenRouter.')
    run.font.name = 'Times New Roman'; run.font.size = Pt(12)
    add_page_break(doc)

    # ── LIST OF FIGURES ───────────────────────────────────────
    centered(doc, 'List of Figures', 14, bold=True)
    doc.add_paragraph()
    figures = [
        ('Figure 1', "Russell's Circumplex Model of Affect and Ekman's Basic Emotions", 'XX'),
        ('Figure 2', 'Existing emotional self-care applications — comparative overview', 'XX'),
        ('Figure 3', 'Mind Space — Static Context Diagram', 'XX'),
        ('Figure 4', 'Mind Space — Package Diagram', 'XX'),
        ('Figure 5', 'Use Case Diagram — Guest User package', 'XX'),
        ('Figure 6', 'Use Case Diagram — Authenticated User package', 'XX'),
        ('Figure 7', 'Use Case Diagram — Sage AI package', 'XX'),
        ('Figure 8', 'Analysis Class Diagram', 'XX'),
        ('Figure 9', 'Design Class Diagram', 'XX'),
        ('Figure 10', 'Sequence Diagram SD1 — User Authentication', 'XX'),
        ('Figure 11', 'Sequence Diagram SD2 — Send Message with Safety Check', 'XX'),
        ('Figure 12', 'Sequence Diagram SD3 — Chat Streaming via SSE', 'XX'),
        ('Figure 13', 'Sequence Diagram SD4 — Wrap Up: Reflection Generation and Arc Assignment', 'XX'),
        ('Figure 14', 'Sequence Diagram SD5 — Generate Arc Insight', 'XX'),
        ('Figure 15', 'Activity Diagram — Arc Assignment Algorithm', 'XX'),
        ('Figure 16', 'Component Diagram', 'XX'),
        ('Figure 17', 'Deployment Diagram', 'XX'),
        ('Figure 18', 'Onboarding screens', 'XX'),
        ('Figure 19', 'Authentication screen', 'XX'),
        ('Figure 20', 'Home screen', 'XX'),
        ('Figure 21', 'Chat screen — Free mode (Sage AI conversation)', 'XX'),
        ('Figure 22', 'Chat screen — Arc mode with carry-forward card', 'XX'),
        ('Figure 23', 'Post-session reflection card', 'XX'),
        ('Figure 24', 'History — Timeline view and Arcs grid', 'XX'),
        ('Figure 25', 'Arc Detail screen with emotion journey graph', 'XX'),
        ('Figure 26', 'Arc Insight screen (Arc Analysis)', 'XX'),
        ('Figure 27', 'Crisis safety screen', 'XX'),
        ('Figure 28', 'Profile and Settings screen', 'XX'),
    ]
    fig_table = doc.add_table(rows=len(figures)+1, cols=3)
    fig_table.style = 'Table Grid'
    for cell, txt in zip(fig_table.rows[0].cells, ['Figure', 'Title', 'Page']):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (num, title, page) in enumerate(figures):
        row = fig_table.rows[i+1].cells
        for cell, txt in zip(row, [num, title, page]):
            cell.text = txt
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    add_page_break(doc)

    # ── LIST OF TABLES ────────────────────────────────────────
    centered(doc, 'List of Tables', 14, bold=True)
    doc.add_paragraph()
    tables_list = [
        ('Table 1', 'List of Abbreviations', 'XX'),
        ('Table 2', 'Comparison of existing emotional self-care applications', 'XX'),
        ('Table 3', 'Functional Requirements (FR)', 'XX'),
        ('Table 4', 'Non-Functional Requirements (NFR)', 'XX'),
        ('Table 5', 'Use Case UC1 — Textual description: User Authentication', 'XX'),
        ('Table 6', 'Use Case UC2 — Textual description: Start Free Conversation', 'XX'),
        ('Table 7', 'Use Case UC3 — Textual description: Send Message with Safety Check', 'XX'),
        ('Table 8', 'Use Case UC4 — Textual description: Chat Streaming via SSE', 'XX'),
        ('Table 9', 'Use Case UC5 — Textual description: Wrap Up Session', 'XX'),
        ('Table 10', 'Use Case UC6 — Textual description: Generate Arc Insight', 'XX'),
        ('Table 11', 'SD1 — Textual description: User Authentication', 'XX'),
        ('Table 12', 'SD2 — Textual description: Send Message with Safety Check', 'XX'),
        ('Table 13', 'SD3 — Textual description: Chat Streaming via SSE', 'XX'),
        ('Table 14', 'SD4 — Textual description: Wrap Up and Arc Assignment', 'XX'),
        ('Table 15', 'SD5 — Textual description: Generate Arc Insight', 'XX'),
        ('Table 16', 'Database schema — profiles table', 'XX'),
        ('Table 17', 'Database schema — chats table', 'XX'),
        ('Table 18', 'Database schema — messages table', 'XX'),
        ('Table 19', 'Database schema — reflections table', 'XX'),
        ('Table 20', 'Database schema — arcs table', 'XX'),
        ('Table 21', 'Database schema — arc_insights table', 'XX'),
        ('Table 22', 'EFT processing stages and Sage behaviour', 'XX'),
    ]
    tbl_tbl = doc.add_table(rows=len(tables_list)+1, cols=3)
    tbl_tbl.style = 'Table Grid'
    for cell, txt in zip(tbl_tbl.rows[0].cells, ['Table', 'Title', 'Page']):
        cell.text = txt
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)
    for i, (num, title, page) in enumerate(tables_list):
        row = tbl_tbl.rows[i+1].cells
        for cell, txt in zip(row, [num, title, page]):
            cell.text = txt
            cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    add_page_break(doc)

    # ── LIST OF ABBREVIATIONS ─────────────────────────────────
    centered(doc, 'List of Abbreviations', 14, bold=True)
    doc.add_paragraph()
    abbrs = [
        ('AI',       'Artificial Intelligence'),
        ('API',      'Application Programming Interface'),
        ('Arc',      'Automatic Recurring Cluster (Mind Space-specific term)'),
        ('CBT',      'Cognitive Behavioural Therapy'),
        ('CTA',      'Call To Action'),
        ('DB',       'Database'),
        ('EFT',      'Emotion-Focused Therapy'),
        ('FR',       'Functional Requirement'),
        ('GDPR',     'General Data Protection Regulation'),
        ('HTTP',     'Hypertext Transfer Protocol'),
        ('IEEE',     'Institute of Electrical and Electronics Engineers'),
        ('JSON',     'JavaScript Object Notation'),
        ('JWT',      'JSON Web Token'),
        ('LLM',      'Large Language Model'),
        ('mHealth',  'Mobile Health'),
        ('ML',       'Machine Learning'),
        ('NFR',      'Non-Functional Requirement'),
        ('NLP',      'Natural Language Processing'),
        ('OAuth',    'Open Authorisation'),
        ('REST',     'Representational State Transfer'),
        ('RLS',      'Row-Level Security'),
        ('RUP',      'Rational Unified Process'),
        ('SD',       'Sequence Diagram'),
        ('SRS',      'Software Requirements Specification'),
        ('SSE',      'Server-Sent Events'),
        ('UC',       'Use Case'),
        ('UI',       'User Interface'),
        ('UML',      'Unified Modelling Language'),
        ('UUID',     'Universally Unique Identifier'),
        ('WHO',      'World Health Organisation'),
    ]
    abbr_table(doc, abbrs)
    add_page_break(doc)

    # ── TABLE OF CONTENTS ─────────────────────────────────────
    centered(doc, 'Table of Contents', 14, bold=True)
    doc.add_paragraph()
    toc_entries = [
        ('General Introduction', ''),
        ('Chapter 1 – Project Specification and Requirements', ''),
        ('  1.1 Introduction', ''),
        ('  1.2 Context and Problem Statement', ''),
        ('  1.3 Key Definitions', ''),
        ('  1.4 Psychological Foundations', ''),
        ('  1.5 Analysis of Existing Solutions', ''),
        ('  1.6 Proposed Solution — Mind Space', ''),
        ('  1.7 Software Requirements Specification (IEEE 830)', ''),
        ('  1.8 Conclusion', ''),
        ('Chapter 2 – System Modelling', ''),
        ('  2.1 Introduction', ''),
        ('  2.2 UML Overview', ''),
        ('  2.3 The RUP Methodology', ''),
        ('  2.4 Use Case Model', ''),
        ('  2.5 Analysis Model', ''),
        ('  2.6 Design Model', ''),
        ('  2.7 Activity Model', ''),
        ('  2.8 Implementation Model', ''),
        ('  2.9 Deployment Model', ''),
        ('  2.10 Conclusion', ''),
        ('Chapter 3 – Implementation', ''),
        ('  3.1 Introduction', ''),
        ('  3.2 Development Environment and Tools', ''),
        ('  3.3 Application Presentation', ''),
        ('  3.4 Conclusion', ''),
        ('General Conclusion', ''),
        ('Bibliography', ''),
    ]
    for entry, page in toc_entries:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(entry)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        if not entry.startswith('  '):
            run.font.bold = True
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # GENERAL INTRODUCTION
    # ══════════════════════════════════════════════════════════
    p = doc.add_paragraph(style='Titre 1 sans numéro')
    p.add_run('General Introduction').font.name = 'Times New Roman'

    body(doc,
         'The modern world has produced a paradox: our ability to connect globally has not '
         'translated into better connection with ourselves. Mental health disorders affect '
         'approximately one billion people worldwide, yet fewer than one in three of those '
         'affected receive any form of care. The barriers are well documented — cost, stigma, '
         'geographic inaccessibility, and long waiting lists for professional support. At the '
         'same time, the smartphone has become the most intimate object in billions of lives, '
         'carried always, opened dozens of times a day, and trusted with our most private '
         'communications.')
    body(doc,
         'This convergence creates both an opportunity and a responsibility. Mobile applications '
         'can reach people where clinical services cannot. But the promise of technology-assisted '
         'emotional support has, to date, been only partially fulfilled. Existing tools tend to '
         'fall into one of two categories: rigid, protocol-based cognitive-behavioural therapy '
         '(CBT) bots that feel mechanical, or passive journalling tools that aggregate mood data '
         'without producing insight. Neither provides the sense of being genuinely heard, which '
         'research consistently identifies as the most important factor in emotional processing.')
    body(doc,
         'This project proposes Mind Space, a mobile application for guided emotional '
         'self-reflection powered by artificial intelligence. Mind Space takes a different '
         'approach: instead of prescribing exercises or tracking mood scores, it creates a space '
         'for open-ended conversation with an AI companion named Sage, generates post-session '
         'reflections informed by Emotion-Focused Therapy (EFT) principles, and — over time — '
         "reveals the recurring emotional themes in a user's life through an automatic clustering "
         'mechanism called the Arc system.')
    body(doc,
         'The technical contribution of this work is threefold. First, a complete Android '
         'application is implemented using Flutter with a Riverpod state architecture and a '
         'fully serverless backend on Supabase. Second, an EFT-informed prompt engineering '
         'methodology is developed and evaluated for non-clinical conversational AI. Third, '
         'an embedding-based automatic thematic detection system is implemented, in which '
         'post-session reflections are encoded as 1536-dimensional semantic vectors and '
         'clustered using cosine similarity with a calibrated threshold to form evolving '
         'emotional arcs.')
    body(doc,
         'This report is organised into three chapters. Chapter 1 establishes the context '
         'and requirements: the global mental health challenge, the psychological and AI '
         'foundations of the work, an analysis of existing solutions, and a full Software '
         'Requirements Specification following the IEEE 830-1998 standard. Chapter 2 presents '
         'the system modelling using the Rational Unified Process (RUP) methodology, including '
         'all six required UML model types. Chapter 3 documents the implementation: the '
         'technology stack, the development environment, and a screen-by-screen presentation '
         'of the completed application.')
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 1
    # ══════════════════════════════════════════════════════════
    heading1(doc, 'Chapter 1 – Project Specification and Requirements')

    # 1.1 Introduction
    heading2(doc, '1.1  Introduction')
    body(doc,
         'This chapter establishes the foundation of the Mind Space project. We begin by '
         'contextualising the global mental health challenge and the role that mobile technology '
         'can play in addressing it. We then review the key definitions, the psychological '
         'frameworks that informed the design, and an analysis of existing solutions. The chapter '
         'concludes with the proposed solution and a complete Software Requirements Specification '
         'structured according to the IEEE 830-1998 standard.')

    # 1.2 Context
    heading2(doc, '1.2  Context and Problem Statement')
    heading3(doc, '1.2.1  The Global Mental Health Challenge')
    body(doc,
         'According to the World Health Organisation (WHO), one in every eight people globally '
         'lives with a mental health disorder. Depression and anxiety disorders alone affect '
         'approximately 280 million and 301 million people respectively. The treatment gap — '
         'the proportion of people with mental health conditions who receive no treatment — '
         'remains above 70% in low- and middle-income countries, and above 50% even in '
         'high-income countries. The economic burden is equally significant: mental health '
         'conditions account for an estimated $1 trillion in lost productivity globally each year.')
    body(doc,
         'The barriers to accessing mental health care are structural. Professional psychotherapy '
         'is expensive, often costing between $100 and $250 per session in private markets. '
         'Waiting times in public health systems can extend to months or years. Geographic '
         'concentration of mental health professionals in urban centres leaves rural populations '
         'with few options. Social stigma continues to prevent individuals from seeking help, '
         'particularly in cultural contexts where emotional vulnerability is perceived negatively. '
         'In Algeria specifically, where this application is primarily targeted, there is a ratio '
         'of fewer than 1.5 psychiatrists per 100,000 population, well below the WHO recommended '
         'level of 3 per 100,000.')

    heading3(doc, '1.2.2  Technological Approaches to Emotional Support')
    body(doc,
         'The rapid proliferation of smartphones — with over 6.5 billion users worldwide — has '
         'made mobile health (mHealth) applications a promising vector for reaching underserved '
         'populations. Mental health applications have grown rapidly, with the global mental '
         'health app market projected to exceed $17 billion by 2030. However, this growth has '
         'not been matched by evidence of efficacy or trustworthiness. Many commercial mental '
         'health applications share sensitive emotional data with third-party advertising '
         'networks, lack clinical grounding, and fail to sustain engagement beyond the first '
         'week of use.')
    body(doc,
         'Recent advances in large language models (LLMs) have introduced a new category of '
         'possibility: conversational AI companions capable of empathic, contextually aware '
         'dialogue. Models such as LLaMA 3 (Meta) and GPT-4 (OpenAI) have demonstrated '
         'sufficient natural language understanding to sustain meaningful conversation in '
         'emotional contexts. The challenge lies in deploying these capabilities in a manner '
         'that is safe, psychologically grounded, and respectful of user privacy.')

    heading3(doc, '1.2.3  Problem Statement')
    body(doc,
         'Existing mobile mental health tools fail to provide the combination of genuine empathic '
         'presence, longitudinal insight, and privacy-first architecture that emotionally '
         'vulnerable users require. Rigid CBT-protocol bots feel mechanical; passive mood '
         'trackers produce data without meaning; journalling apps lack structure and reflection. '
         'None of the major commercial offerings automatically detect and surface the recurring '
         'emotional themes that would give users a broader view of their inner life. This project '
         'addresses this gap.')

    # 1.3 Definitions
    heading2(doc, '1.3  Key Definitions')
    heading3(doc, '1.3.1  Mental Health')
    body(doc,
         'The WHO defines mental health as "a state of mental well-being that enables people '
         'to cope with the stresses of life, realise their abilities, learn well and work well, '
         'and contribute to their community." It is not merely the absence of mental disorders '
         'but encompasses emotional, psychological, and social well-being. For the purposes of '
         'this project, we adopt the broader definition and focus specifically on emotional '
         "self-awareness and resilience — the capacity to recognise, name, and process one's "
         'own emotional states.')
    heading3(doc, '1.3.2  Mobile Health (mHealth)')
    body(doc,
         'mHealth refers to the use of mobile devices, wearable sensors, and wireless '
         'communication technology to deliver health services and information. In the mental '
         'health domain, mHealth applications range from symptom trackers and CBT exercise '
         'platforms to AI-powered conversational companions. The defining characteristic of '
         'mHealth is ubiquitous accessibility: unlike clinical services, a mobile application '
         'is available at any moment of the day, without appointment, and in environments '
         'where professional help is unavailable.')
    heading3(doc, '1.3.3  Artificial Intelligence in Mental Health')
    body(doc,
         'Artificial intelligence, in the context of mental health applications, encompasses '
         'natural language processing for conversational agents, machine learning models for '
         'emotion detection, and deep learning systems for clinical risk assessment. '
         'Conversational AI agents have been shown to reduce symptoms of depression and anxiety '
         'in controlled studies when combined with structured therapeutic frameworks such as CBT '
         'or mindfulness. However, the use of AI in mental health contexts raises significant '
         'ethical questions around data sensitivity, model hallucination, and the risk of '
         'inappropriate responses to crisis situations.')
    heading3(doc, '1.3.4  Large Language Models (LLMs)')
    body(doc,
         'Large Language Models are neural network architectures trained on massive corpora of '
         'text to predict the most probable continuation of a given sequence. Contemporary LLMs '
         'such as LLaMA 3.3 70B (Meta, accessed via the Groq API) demonstrate emergent '
         'capabilities including in-context learning, multi-step reasoning, and stylistic '
         'adaptation. In Mind Space, LLMs serve three distinct roles: conversational agent '
         '(Sage AI), safety classifier, and longitudinal insight generator. Each role uses a '
         'different model and temperature setting, optimised for its specific task.')
    heading3(doc, '1.3.5  Semantic Embeddings')
    body(doc,
         'A semantic embedding is a dense vector representation of a piece of text, produced '
         'by an encoder model trained to place semantically similar texts close together in '
         'vector space. In this project, the OpenAI text-embedding-3-small model is used to '
         'produce 1536-dimensional embeddings of post-session reflection summaries. These '
         'embeddings enable automatic thematic clustering of sessions using cosine similarity, '
         'without requiring any manual categorisation from the user.')

    # 1.4 Psychological Foundations
    heading2(doc, '1.4  Psychological Foundations')
    heading3(doc, '1.4.1  Expressive Writing and Emotional Processing')
    body(doc,
         'The therapeutic value of expressive writing was established by James Pennebaker and '
         'colleagues in a series of landmark studies beginning in the 1980s. Pennebaker\'s '
         'research demonstrated that structured written disclosure of emotionally significant '
         'events produces measurable improvements in psychological and physical health outcomes, '
         'including reduced physician visits, improved immune function, and lower self-reported '
         'distress. The mechanism proposed is that expressive writing facilitates cognitive '
         'processing — the construction of a narrative around difficult experiences — which '
         'reduces the cognitive load associated with thought suppression. Mind Space applies '
         'this principle by providing a structured, low-friction space for daily emotional '
         'disclosure through conversation.')
    heading3(doc, '1.4.2  Emotion-Focused Therapy (EFT)')
    body(doc,
         'Emotion-Focused Therapy, developed by Leslie Greenberg and colleagues, is a '
         'psychotherapeutic approach that prioritises the identification, expression, and '
         'transformation of emotional experience. A central concept in EFT is the distinction '
         'between primary emotions (the core emotional response to a situation), secondary '
         'emotions (surface-level reactions that mask the primary emotion), and instrumental '
         'emotions (emotions expressed to achieve a social goal). Effective emotional processing '
         'requires accessing the primary emotion rather than being deflected by the secondary. '
         'In Mind Space, the Sage AI is instructed to reflect back the layer beneath the surface '
         'emotion expressed by the user, guided by an EFT-informed prompt. Each post-session '
         'reflection captures what_sage_heard — the primary emotional truth of the session, '
         'not the surface label.')
    body(doc,
         'EFT also introduces the concept of processing stages, through which a client moves '
         'from initial awareness and expression of a problem through to emotional regulation '
         'and eventual integration. Mind Space implements a six-stage EFT processing model '
         'per Arc — forming, venting, stabilising, processing, shifting, and integrating — '
         'which modulates the conversational strategy of Sage across sessions.')
    heading3(doc, '1.4.3  Emotion Classification Models')
    body(doc,
         'Two foundational models of emotion structure inform the design of the emotion system '
         'in Mind Space. Paul Ekman identified six basic universal emotions — happiness, sadness, '
         'fear, disgust, anger, and surprise — expressed consistently across cultures through '
         'facial configurations. James Russell\'s Circumplex Model of Affect organises emotional '
         'states along two continuous dimensions: valence (positive to negative) and arousal '
         '(activated to deactivated). Mind Space does not use Ekman\'s basic categories directly, '
         'as they are too coarse for nuanced emotional reflection. Instead, a set of six '
         'contextually richer emotional spirit archetypes is defined — anxious, calm, frustrated, '
         'sad, hopeful, and overwhelmed — each mapped to a visual representation, colour, and '
         'animation style within the application.')

    add_figure(doc, img('page14_img00.png'),
               'Figure 1 – Russell\'s Circumplex Model of Affect (left) and Ekman\'s basic emotions (right)',
               width_cm=13)

    # 1.5 Existing Solutions
    heading2(doc, '1.5  Analysis of Existing Solutions')
    body(doc,
         'A review of the leading mental health and emotional support applications currently '
         'available on the market reveals several notable limitations. Table 2 summarises '
         'the key criteria against which each solution was evaluated.')

    add_figure(doc, img('page16_img00.jpeg'),
               'Figure 2 – Comparative overview of existing emotional self-care applications',
               width_cm=13)

    heading3(doc, '1.5.1  Wysa')
    body(doc,
         'Wysa is an AI-powered emotional wellbeing assistant that combines a conversational '
         'interface with structured CBT, dialectical behaviour therapy (DBT), and mindfulness '
         'exercises. It is among the most clinically validated consumer mental health applications, '
         'with published studies demonstrating symptom reduction for mild-to-moderate depression '
         'and anxiety. However, Wysa\'s conversations are guided by decision-tree scripts rather '
         'than large language models, producing interactions that feel formulaic over time. '
         'Privacy analysis has raised concerns about data sharing with analytics platforms.')
    heading3(doc, '1.5.2  Woebot')
    body(doc,
         'Woebot is a CBT-based conversational agent developed by researchers from Stanford '
         'University. It demonstrated significant reductions in depression symptoms in a '
         'randomised controlled trial published in 2017, establishing its clinical credibility. '
         'Like Wysa, Woebot uses rule-based dialogue management rather than LLMs, which limits '
         'its expressiveness and adaptability. Its terms of service have historically permitted '
         'use of anonymised conversation data for research and product improvement, which '
         'raises privacy concerns for sensitive emotional content.')
    heading3(doc, '1.5.3  Daylio')
    body(doc,
         'Daylio is a micro-journalling and mood tracking application that focuses on quantitative '
         'data capture rather than reflective conversation. Users log moods and activities from '
         'predefined lists, producing statistical summaries of emotional patterns over time. '
         'While effective for tracking, Daylio provides no qualitative insight into the meaning '
         'behind the patterns it surfaces. It has no conversational component and does not '
         'provide safety safeguards for users in distress.')
    heading3(doc, '1.5.4  Comparative Analysis')
    comparison_table(doc)
    body(doc,
         'Table 2 illustrates the key differentiating features of Mind Space relative to '
         'existing solutions. The most significant gaps in the market are: the absence of '
         'EFT-informed conversational depth, the lack of automatic longitudinal theme detection, '
         'and inadequate crisis safety architecture. Mind Space addresses all three.')

    # 1.6 Proposed Solution
    heading2(doc, '1.6  Proposed Solution — Mind Space')
    heading3(doc, '1.6.1  Vision and Core Concept')
    body(doc,
         'Mind Space is designed around a single insight: the most valuable emotional support '
         'is not instructions or tracking, but the experience of being genuinely heard. The '
         'application creates a private, persistent space for open-ended conversation with an '
         'AI companion named Sage — a presence trained to listen with empathy, reflect back '
         'the emotional truth beneath the surface, and gently help the user understand their '
         'own patterns over time.')
    body(doc,
         'The user interface is designed to feel like a personal, intimate tool — not a '
         'clinical application. The visual language draws on deep-space imagery and organic '
         'forms: a near-black purple environment, glowing ambient orbs representing emotional '
         'states, and gentle animation throughout. The experience is intentionally slow and '
         'contemplative, contrasting with the speed and stimulation of most mobile applications.')
    heading3(doc, '1.6.2  The Arc System')
    body(doc,
         'The Arc system is the core technical contribution of this work. An Arc is a '
         'persistent, evolving cluster of conversation sessions that share a common emotional '
         'theme — for example, "The Job Hunt", "My Relationship with Dad", or "Feeling Stuck". '
         'Arcs emerge automatically, without the user needing to name or categorise sessions. '
         'At the end of each session, the system generates a structured reflection and encodes '
         'it as a semantic embedding. This embedding is compared to the centroids of all '
         'existing Arcs. If the similarity exceeds a calibrated threshold of 0.78, the '
         'session is assigned to the matching Arc; if not, a new Arc is created. Over time, '
         'each Arc accumulates sessions and can generate a macro-level AI insight — a '
         'synthesised analysis of how the user\'s relationship with that theme has evolved.')
    heading3(doc, '1.6.3  The Safety Architecture')
    body(doc,
         'A dedicated safety classifier operates as a pre-processing gate on every user '
         'message before it reaches the main language model. This classifier uses a smaller, '
         'faster model (LLaMA 3.1 8B instant, via Groq) at near-zero temperature (0.1) to '
         'produce a binary SAFE/CRISIS classification. The system is designed to fail closed: '
         'if the classification API is unavailable or times out, the system defaults to the '
         'CRISIS state and presents pre-written emergency resources — including the Algerian '
         'emergency number SAMU 15 — rather than allowing the message to proceed to the '
         'general model. Messages classified as CRISIS are saved to the database with a '
         'flagged_for_safety flag and never forwarded to any external service beyond the '
         'safety classifier itself.')

    # 1.7 SRS
    heading2(doc, '1.7  Software Requirements Specification (IEEE 830-1998)')
    heading3(doc, '1.7.1  Purpose')
    body(doc,
         'This Software Requirements Specification (SRS) defines the functional and '
         'non-functional requirements for Mind Space, version 1.0 (Android). It is intended '
         'for the development team, the project supervisor, and any evaluator assessing '
         'the completeness and correctness of the system. This document follows the structure '
         'prescribed by IEEE Standard 830-1998.')
    heading3(doc, '1.7.2  Scope')
    body(doc,
         'Mind Space is a mobile application providing AI-assisted emotional self-reflection '
         'for Android users. The system allows users to: initiate open-ended or theme-guided '
         'conversations with an AI companion; receive structured post-session reflections '
         'informed by EFT principles; observe and revisit their conversation history; and '
         'track the evolution of recurring emotional themes through the Arc system. The '
         'application does not provide clinical diagnosis, medical advice, or human therapist '
         'interaction. It is not intended as a substitute for professional mental health care.')
    heading3(doc, '1.7.3  Definitions, Acronyms, Abbreviations')
    body(doc,
         'All abbreviations used in this document are defined in the List of Abbreviations '
         'at the front of this report. Key project-specific terms: '
         'Arc — a persistent cluster of emotionally related sessions; '
         'Sage — the AI conversational companion within Mind Space; '
         'Reflection — a structured post-session summary generated by the AI; '
         'Spirit — a visual/conceptual archetype representing a primary emotional state '
         '(anxious, calm, frustrated, sad, hopeful, overwhelmed); '
         'Wrap Up — the user action that ends a session and triggers reflection generation.')
    heading3(doc, '1.7.4  References')
    body(doc,
         '[1] IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements '
         'Specifications. IEEE, 1998. '
         '[2] Greenberg, L. S. (2004). Emotion-focused therapy. Clinical Psychology and '
         'Psychotherapy, 11(1), 3–16. '
         '[3] Pennebaker, J. W. (1997). Writing about emotional experiences as a therapeutic '
         'process. Psychological Science, 8(3), 162–166. '
         '[4] World Health Organisation (2022). World Mental Health Report. Geneva: WHO. '
         '[5] Kruchten, P. (2003). The Rational Unified Process: An Introduction (3rd ed.). '
         'Addison-Wesley. '
         '[6] Booch, G., Rumbaugh, J., & Jacobson, I. (2005). The Unified Modeling Language '
         'User Guide (2nd ed.). Addison-Wesley.')

    heading3(doc, '1.7.5  General Description')
    heading4(doc, '1.7.5.1  Product Perspective')
    body(doc,
         'Mind Space is a standalone mobile application for Android. It does not replace '
         'any existing system; it introduces a new product category for the targeted user '
         'population. The system communicates with three external service providers via '
         'HTTPS: Supabase (authentication, database, edge functions), Groq (LLM inference), '
         'and OpenAI (text embeddings). A fourth service, OpenRouter, is used for '
         'long-context arc insight generation. All external communications are encrypted '
         'in transit. No emotional data is stored on the device beyond the active session; '
         'all persistent data resides in the PostgreSQL database hosted on Supabase Cloud.')
    heading4(doc, '1.7.5.2  Product Functions')
    body(doc, 'The major functions of Mind Space are:', indent=False)
    for f in [
        'Guided conversation with the Sage AI companion (streaming, token-by-token).',
        'Pre-message safety classification (SAFE/CRISIS gate before every message).',
        'Post-session reflection generation: spirit, what_sage_heard, question_to_sit_with, shared_perspective.',
        'Automatic arc assignment via embedding cosine similarity (pgvector, threshold 0.78).',
        'Session history browsing — Timeline (reverse-chronological) and Arcs grid views.',
        'Arc detail view: session list, EFT processing stage, emotion journey chart.',
        'Arc insight generation: longitudinal analysis via DeepSeek R1 (unlocks at ≥3 sessions).',
        'User account management: Google OAuth, email magic link, export data, delete account.',
    ]:
        bullet(doc, f)
    doc.add_paragraph()

    heading4(doc, '1.7.5.3  User Characteristics')
    body(doc,
         'The primary user is an adult (18+) smartphone user experiencing emotional difficulty '
         'who is seeking a private, accessible space for self-reflection. The user may have '
         'no formal mental health background. Technical literacy is assumed to be at the level '
         'of a regular smartphone user; no specialised knowledge is required. The secondary '
         'user type is someone with an existing reflective practice (journalling, therapy) '
         'who seeks a structured AI-assisted supplement to that practice.')
    heading4(doc, '1.7.5.4  Constraints')
    body(doc,
         'The application requires an active internet connection for all AI-dependent features '
         '(conversation, reflection generation, arc assignment, arc insight). The safety '
         'classifier is the most critical real-time constraint: it must respond within 3 seconds '
         'or the system defaults to CRISIS state. The application is limited to Android for '
         'version 1.0. All AI operations depend on third-party API availability; service '
         'outages will degrade functionality. The application does not store embeddings or '
         'AI responses locally; offline use is not supported. User data is subject to the '
         'Supabase data processing terms and stored in the EU-West region.')
    heading4(doc, '1.7.5.5  Assumptions and Dependencies')
    body(doc,
         'This SRS assumes that: (1) the user owns an Android 6.0 (API 23) or higher device; '
         '(2) the user has a stable internet connection; (3) the Groq, OpenAI, and OpenRouter '
         'free-tier API limits are sufficient for a development and limited-deployment context '
         '(Groq: 1,000+ req/day on llama-3.3-70b; OpenAI: pay-per-use; OpenRouter: 200 RPD on '
         'deepseek-r1:free); (4) Supabase remains in the free tier for the scope of this project. '
         'The system depends on the continued availability of these external APIs. Any '
         'changes to their interfaces or pricing models may require adaptation of the '
         'corresponding edge functions.')

    heading3(doc, '1.7.6  Specific Requirements')
    heading4(doc, '1.7.6.1  Functional Requirements')
    fr_rows = [
        ('FR-01', 'The system shall allow a guest user to view onboarding slides (3 screens) and proceed to authentication without prior account creation.', 'Essential'),
        ('FR-02', 'The system shall authenticate users via Google OAuth 2.0 or email magic link through Supabase Auth and create a profile record on first login.', 'Essential'),
        ('FR-03', 'The system shall allow an authenticated user to start a new open-ended conversation session with Sage at any time from the home screen.', 'Essential'),
        ('FR-04', 'The system shall classify every user message as SAFE or CRISIS before processing; CRISIS messages shall display pre-written resources and not be forwarded to the main LLM.', 'Essential'),
        ('FR-05', 'The system shall stream Sage\'s response token-by-token via SSE, updating the UI in real time without page reload.', 'Essential'),
        ('FR-06', 'The system shall generate a structured EFT reflection (spirit, what_sage_heard, question_to_sit_with, shared_perspective) upon Wrap Up, using LLaMA 3.3 70B via Groq.', 'Essential'),
        ('FR-07', 'The system shall encode each reflection as a 1536-dimensional embedding and assign it to the closest Arc or create a new Arc using cosine similarity (threshold 0.78).', 'Essential'),
        ('FR-08', 'The system shall allow the user to browse past sessions in a reverse-chronological Timeline view and in an Arcs grid view, with filter by emotion.', 'Essential'),
        ('FR-09', 'The system shall display an Arc detail page showing all sessions in the Arc, the EFT processing stage, an emotion journey chart, and a Generate Insight CTA (visible when session_count ≥ 3).', 'Essential'),
        ('FR-10', 'The system shall generate a longitudinal Arc insight (how_it_evolved, pattern_noticed) using DeepSeek R1 via OpenRouter, cached for 7 days or until a new session is added.', 'Important'),
        ('FR-11', 'The system shall allow the user to export all their data as a JSON file and delete their account, including all associated data, from the Settings screen.', 'Important'),
        ('FR-12', 'The system shall present a Wrap Up action after the user has sent 4 or more messages in a session, and allow wrapping up at any point after that.', 'Important'),
    ]
    body(doc, 'Table 3 – Functional Requirements', bold=True, indent=False)
    requirement_table(doc, fr_rows)

    heading4(doc, '1.7.6.2  Non-Functional Requirements')
    nfr_rows = [
        ('NFR-01', 'Performance: The application shall load the home screen from a cold start in under 2 seconds on a mid-range Android device (3 GB RAM, average 4G connection).', 'Essential'),
        ('NFR-02', 'Safety: The safety classifier shall have zero false negatives (no CRISIS message must ever pass undetected); the system shall fail closed on classifier API failure.', 'Critical'),
        ('NFR-03', 'Privacy: No emotional data (messages, reflections, embeddings) shall be transmitted to third-party advertising or analytics services. Data at rest is protected by Supabase Row-Level Security (RLS) policies.', 'Essential'),
        ('NFR-04', 'Reliability: The system shall maintain chat functionality (degraded: no streaming, static reply) if the SSE connection is interrupted, and display an offline banner to the user.', 'Important'),
        ('NFR-05', 'Usability: A new user shall be able to complete the onboarding flow and start a first conversation within 90 seconds of installing the application.', 'Important'),
        ('NFR-06', 'Scalability: The Arc assignment algorithm shall complete within 1 second for a user with up to 50 active arcs (pgvector ivfflat index, lists=100).', 'Important'),
    ]
    body(doc, 'Table 4 – Non-Functional Requirements', bold=True, indent=False)
    requirement_table(doc, nfr_rows)

    # 1.8 Conclusion
    heading2(doc, '1.8  Conclusion')
    body(doc,
         'This chapter has established the full context and requirements for the Mind Space '
         'project. The global mental health challenge provides a clear motivation; the '
         'psychological foundations of EFT and expressive writing provide the therapeutic '
         'grounding; and the analysis of existing solutions demonstrates the gap that Mind '
         'Space is designed to fill. The Software Requirements Specification provides a '
         'complete, verifiable set of functional and non-functional requirements following '
         'the IEEE 830-1998 standard. Chapter 2 will translate these requirements into a '
         'formal system model using the RUP methodology and UML notation.')
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 2
    # ══════════════════════════════════════════════════════════
    heading1(doc, 'Chapter 2 – System Modelling')

    # 2.1 Introduction
    heading2(doc, '2.1  Introduction')
    body(doc,
         'This chapter presents the formal system model for Mind Space, developed using the '
         'Rational Unified Process (RUP) methodology and expressed in the Unified Modelling '
         'Language (UML). Six model types are produced: the Use Case model, the Analysis '
         'model, the Design model, the Activity model, the Implementation model, and the '
         'Deployment model. Together, these models provide a complete, multi-perspective '
         'description of the system — from stakeholder requirements through to the physical '
         'infrastructure on which it runs.')

    # 2.2 UML
    heading2(doc, '2.2  UML Overview')
    body(doc,
         'The Unified Modelling Language (UML) is a standardised, general-purpose modelling '
         'language in the field of software engineering, maintained by the Object Management '
         'Group (OMG). UML provides a set of graphical notation techniques for creating '
         'visual models of software systems. It encompasses thirteen diagram types organised '
         'into two categories: structural diagrams (class, component, deployment, object, '
         'package, profile) and behavioural diagrams (activity, communication, interaction '
         'overview, sequence, state machine, timing, use case). In this project, we use '
         'use case diagrams, class diagrams, sequence diagrams, activity diagrams, component '
         'diagrams, and deployment diagrams — corresponding to the six model types required '
         'by the RUP methodology.')

    # 2.3 RUP
    heading2(doc, '2.3  The RUP Methodology')
    heading3(doc, '2.3.1  RUP Phases')
    body(doc,
         'The Rational Unified Process (RUP) is an iterative software development process '
         'framework developed by Rational Software (later acquired by IBM) and described by '
         'Philippe Kruchten. RUP organises development into four phases, each with specific '
         'objectives and milestones:')
    for phase_item in [
        'Inception: Establishes the scope, feasibility, and business case of the project. Produces the initial use case model and risk assessment.',
        'Elaboration: Refines the requirements, produces the architecture baseline, and mitigates the most significant technical risks. Produces the analysis and design models.',
        'Construction: Implements and tests all features against the elaborated architecture. Produces the implementation and test models.',
        'Transition: Delivers the software to the end user. Includes beta testing, user training, and final deployment.',
    ]:
        bullet(doc, phase_item)
    doc.add_paragraph()

    heading3(doc, '2.3.2  RUP Models Applied in This Work')
    body(doc,
         'The six UML models produced in this chapter correspond to the six RUP model types '
         'as follows:')
    for m in [
        'Use Case Model (Inception/Elaboration): captures the functional requirements from the user perspective via static context diagram, package diagram, use case diagrams, and textual descriptions.',
        'Analysis Model (Elaboration): captures the conceptual structure of the system via an analysis class diagram, focusing on entities and their relationships without implementation detail.',
        'Design Model (Elaboration/Construction): refines the analysis model into a design class diagram and sequence diagrams that describe the dynamic behaviour of the system.',
        'Activity Model (Construction): describes the algorithmic flow of the Arc assignment process as an activity diagram.',
        'Implementation Model (Construction): describes the software components and their dependencies via a component diagram.',
        'Deployment Model (Transition): describes the physical nodes and the artefacts deployed on each, via a deployment diagram.',
    ]:
        bullet(doc, m)
    doc.add_paragraph()

    # 2.4 Use Case Model
    heading2(doc, '2.4  Use Case Model')
    heading3(doc, '2.4.1  System Context Diagram')
    body(doc,
         'The static context diagram presents Mind Space as a single system boundary interacting '
         'with four external actors: the Guest User (unauthenticated), the Authenticated User '
         '(signed in), the Sage AI backend subsystem (which initiates LLM inference pipelines), '
         'and the External APIs (Groq, OpenAI, and OpenRouter collectively). The diagram '
         'establishes the scope of the system and the nature of each external interaction.')
    add_figure(doc, diag('fig_context_diagram.png'),
               'Figure 3 – Mind Space: Static Context Diagram', width_cm=13)

    heading3(doc, '2.4.2  Package Diagram')
    body(doc,
         'The package diagram organises the use cases into three packages, each associated '
         'with one primary actor. The Guest User package contains the onboarding and '
         'authentication use cases. The Authenticated User package contains all core '
         'application features — conversation, history, arc management, and settings. The '
         'Sage AI package contains the backend processing use cases triggered by user actions '
         '— safety classification, LLM streaming, reflection generation, embedding, '
         'arc assignment, and insight generation. The Authenticated User package depends on '
         '(«use») the Sage AI package for all AI-driven features.')
    add_figure(doc, diag('fig_package_diagram.png'),
               'Figure 4 – Mind Space: Package Diagram', width_cm=13)

    heading3(doc, '2.4.3  Use Case Diagrams')
    heading4(doc, '2.4.3.1  Guest User Package')
    body(doc,
         'The Guest User package contains two use cases: View Onboarding and Authenticate. '
         'The Guest User actor initiates the application for the first time and is shown '
         'three onboarding slides before being prompted to authenticate. Authentication '
         'includes two alternative flows: Google OAuth and email magic link, both handled '
         'by Supabase Auth. Upon successful authentication, the user transitions to the '
         'Authenticated User role.')
    add_figure(doc, img('page23_img01.png'),
               'Figure 5 – Use Case Diagram: Guest User package', width_cm=11)

    heading4(doc, '2.4.3.2  Authenticated User Package')
    body(doc,
         'The Authenticated User package contains nine use cases covering the full '
         'application feature set: Start Conversation, Send Message, Receive Streaming '
         'Reply, Wrap Up Session, Browse Timeline, Browse Arcs, View Arc Detail, Generate '
         'Arc Insight, and Manage Settings. These use cases span all four screens of the '
         'main navigation (Home, Chat, History/Arcs, Analysis, Settings). The Send Message '
         'use case includes the Sage AI package via a «uses» relationship, as every message '
         'triggers the safety-check and chat-stream edge functions.')
    add_figure(doc, img('page24_img00.png'),
               'Figure 6 – Use Case Diagram: Authenticated User package', width_cm=12)

    heading4(doc, '2.4.3.3  Sage AI Package')
    body(doc,
         'The Sage AI package contains six use cases that represent the backend processing '
         'pipelines: Classify Safety, Stream LLM Reply, Generate Reflection, Embed '
         'Reflection, Assign to Arc, and Generate Arc Insight. These use cases are '
         'initiated by the system in response to Authenticated User actions and involve '
         'calls to the Supabase Edge Functions and external AI service providers. The '
         'Sage AI actor represents the aggregate of these backend AI subsystems.')
    add_figure(doc, img('page23_img00.png'),
               'Figure 7 – Use Case Diagram: Sage AI package', width_cm=10)

    heading3(doc, '2.4.4  Use Case Textual Descriptions')
    body(doc, 'Table 5 – Use Case UC1: User Authentication', bold=True, indent=False)
    uc_table(doc, [
        ('Use Case', 'UC1 — User Authentication'),
        ('Objective', 'Allow a guest user to create an account or sign in to an existing account'),
        ('Primary Actor', 'Guest User'),
        ('Secondary Actor', 'Supabase Auth Service'),
        ('Precondition', 'User has the application installed; device has internet connectivity'),
        ('Postcondition', 'User is authenticated; a profile record exists in the profiles table; onboarding_completed = true after first use'),
        ('Version', '1.0'),
        ('Nominal Scenario',
         '1. User opens the application for the first time.\n'
         '2. System displays three onboarding slides (swipeable).\n'
         '3. User taps "Get Started" on the third slide.\n'
         '4. System displays the authentication screen.\n'
         '5a. (Google) User taps "Sign in with Google" → Google OAuth flow → Supabase validates → JWT issued.\n'
         '5b. (Magic link) User enters email → system sends magic link → user clicks link → Supabase validates → JWT issued.\n'
         '6. Supabase handle_new_user() trigger fires and inserts a row in profiles.\n'
         '7. System navigates to Home screen.'),
        ('Alternative A1', 'Returning user (onboarding_completed = true): step 2–3 are skipped; system navigates directly to Home.'),
        ('Alternative A2', 'User already has an account and signs in: steps 6 is skipped (profile already exists).'),
        ('Error E1', 'Authentication fails (invalid credentials, network error): system displays error message and retry option.'),
        ('Error E2', 'Profile trigger fails: system retries profile read and creates profile if absent.'),
    ])

    body(doc, 'Table 6 – Use Case UC2: Start Free Conversation', bold=True, indent=False)
    uc_table(doc, [
        ('Use Case', 'UC2 — Start Free Conversation'),
        ('Objective', 'Open a new chat session without Arc context'),
        ('Primary Actor', 'Authenticated User'),
        ('Precondition', 'User is authenticated and on the Home screen'),
        ('Postcondition', 'A new chat row is created with status = active; Sage has sent an opening message'),
        ('Version', '1.0'),
        ('Nominal Scenario',
         '1. User taps the ambient orb or "Start a conversation" button.\n'
         '2. System creates a new chat record in the chats table (status = active, arc_id = null).\n'
         '3. Chat screen slides up as a full-screen modal.\n'
         '4. System sends Sage\'s opening message: "What\'s on your mind?".\n'
         '5. Quick reply chips appear (three pre-defined options).\n'
         '6. User begins typing or selects a quick reply.'),
        ('Alternative A1', 'User taps an Arc card: chat opens in Arc mode (arc_id pre-populated; carry-forward card shown above first message).'),
        ('Error E1', 'Network unavailable: system displays offline banner; chat row not created.'),
    ])

    body(doc, 'Table 7 – Use Case UC3: Send Message with Safety Check', bold=True, indent=False)
    uc_table(doc, [
        ('Use Case', 'UC3 — Send Message with Safety Check'),
        ('Objective', 'Classify and process a user message, showing crisis resources if needed'),
        ('Primary Actor', 'Authenticated User'),
        ('Secondary Actor', 'Sage AI (Safety Check pipeline)'),
        ('Precondition', 'User is authenticated; active chat session is open; input field contains non-empty text'),
        ('Postcondition', 'Message is classified and either: processed by LLM (SAFE) or resources shown and message flagged (CRISIS)'),
        ('Version', '1.0'),
        ('Nominal Scenario',
         '1. User types a message and taps the send button.\n'
         '2. System sends POST to safety-check Edge Function.\n'
         '3. Edge function calls Groq llama-3.1-8b-instant (temp 0.1) for SAFE/CRISIS classification.\n'
         '4. SAFE: Edge function returns SAFE flag → system proceeds to chat-stream (UC4).\n'
         '5. Message saved with flagged_for_safety = false.'),
        ('Alternative A1 (CRISIS)', 'Classification returns CRISIS → system shows pre-written resources card (including SAMU 15) → message saved with flagged_for_safety = true → LLM not called.'),
        ('Alternative A2', 'User taps "I\'m safe" on crisis card → chat continues with the next message.'),
        ('Error E1', 'Safety check API is unavailable or times out → system fails CLOSED: crisis card displayed regardless.'),
        ('Error E2', 'Groq rate limit exceeded → treated as classifier failure → system fails CLOSED.'),
    ])

    body(doc, 'Table 8 – Use Case UC4: Chat Streaming via SSE', bold=True, indent=False)
    uc_table(doc, [
        ('Use Case', 'UC4 — Chat Streaming via SSE'),
        ('Objective', 'Stream Sage AI reply token-by-token to the user in real time'),
        ('Primary Actor', 'Authenticated User'),
        ('Secondary Actor', 'Sage AI (chat-stream Edge Function, Groq API)'),
        ('Precondition', 'Message has passed UC3 (SAFE classification); active chat session is open'),
        ('Postcondition', 'Sage reply fully streamed; both user and assistant messages saved to DB; message_count incremented'),
        ('Version', '1.0'),
        ('Nominal Scenario',
         '1. System saves user message (role = user) to the messages table.\n'
         '2. System sends POST to chat-stream Edge Function.\n'
         '3. Edge function fetches the last 20 messages from the chat.\n'
         '4. Edge function calls Groq llama-3.3-70b-versatile with streaming = true, temperature = 0.7.\n'
         '5. Groq streams response tokens via SSE.\n'
         '6. Each token is forwarded to the Flutter client over the open HTTP connection.\n'
         '7. Flutter appends each token to the chat UI in real time.\n'
         '8. Stream completes; full assistant message saved to DB (role = assistant).\n'
         '9. If message_count >= 4: Wrap Up pill becomes visible in the header.'),
        ('Alternative A1', 'User closes the chat mid-stream: partial message is discarded; no partial message saved.'),
        ('Error E1', 'Groq API is unavailable: Edge function returns 503 → app displays offline banner.'),
        ('Error E2', 'SSE connection drops mid-stream: Flutter detects closed stream → offline banner shown; partial buffer discarded.'),
    ])

    body(doc, 'Table 9 – Use Case UC5: Wrap Up Session', bold=True, indent=False)
    uc_table(doc, [
        ('Use Case', 'UC5 — Wrap Up Session'),
        ('Objective', 'End the session, generate an EFT reflection, and assign it to an Arc'),
        ('Primary Actor', 'Authenticated User'),
        ('Secondary Actor', 'Sage AI (end-chat, assign-arc Edge Functions; Groq, OpenAI APIs)'),
        ('Precondition', 'User is authenticated; message_count >= 1; user taps Wrap Up'),
        ('Postcondition', 'Reflection generated and saved; embedding computed; session assigned to an Arc (new or existing); reflection card displayed'),
        ('Version', '1.0'),
        ('Nominal Scenario',
         '1. User taps "Wrap Up" in the chat header or "..." sheet.\n'
         '2. System shows a loading state on the reflection card.\n'
         '3. end-chat Edge Function is called.\n'
         '4. Function fetches the full conversation and the user\'s cross-arc context.\n'
         '5. Groq llama-3.3-70b generates a reflection JSON: {spirit_id, what_sage_heard, question_to_sit_with, shared_perspective}.\n'
         '6. Reflection saved to the reflections table.\n'
         '7. OpenAI text-embedding-3-small encodes what_sage_heard into a 1536-dim vector.\n'
         '8. assign-arc Edge Function compares the vector with all arc centroids using pgvector cosine similarity.\n'
         '9. Reflection assigned to the closest Arc (if similarity >= 0.78) or a new Arc is created (if similarity < 0.65).\n'
         '10. Arc centroid updated; session_count incremented; processing_stage updated.\n'
         '11. Reflection card displayed to user with cascade animation.'),
        ('Alternative A1', 'Existing arc matched (>= 0.78): centroid updated via incremental average formula.'),
        ('Alternative A2', 'Soft zone (0.65–0.78): assigned to closest arc; needs_review = true.'),
        ('Alternative A3', 'New arc created (< 0.65): LLM generates 2–4 word name; dominant spirit assigned.'),
        ('Error E1', 'Reflection generation fails: system shows retry or skip button.'),
        ('Error E2', 'Embedding fails: reflection saved with needs_arc_assignment = true for background retry.'),
    ])

    body(doc, 'Table 10 – Use Case UC6: Generate Arc Insight', bold=True, indent=False)
    uc_table(doc, [
        ('Use Case', 'UC6 — Generate Arc Insight'),
        ('Objective', 'Generate a longitudinal macro-analysis of an Arc with >= 3 sessions'),
        ('Primary Actor', 'Authenticated User'),
        ('Secondary Actor', 'Sage AI (generate-arc-insight Edge Function, OpenRouter/DeepSeek R1)'),
        ('Precondition', 'User is authenticated; Arc has >= 3 sessions; user taps "Generate Arc Analysis"'),
        ('Postcondition', 'Arc insight generated and saved; how_it_evolved and pattern_noticed displayed; result cached 7 days'),
        ('Version', '1.0'),
        ('Nominal Scenario',
         '1. User navigates to Arc Detail and taps "Generate Arc Analysis".\n'
         '2. System shows loading state.\n'
         '3. generate-arc-insight Edge Function checks for a valid cached insight.\n'
         '4. No valid cache: function fetches all arc reflections in chronological order.\n'
         '5. Function includes cross-arc context (other active arcs) in the prompt.\n'
         '6. OpenRouter DeepSeek R1 (temperature 0.3, max 1500 tokens) generates {how_it_evolved, pattern_noticed}.\n'
         '7. Insight saved with session_count_at_generation snapshot.\n'
         '8. Arc Insight screen displayed to user.'),
        ('Alternative A1', 'Valid cache exists (< 7 days AND session_count unchanged): cached insight returned immediately.'),
        ('Error E1', 'OpenRouter unavailable: Edge function returns error → app shows retry button.'),
        ('Error E2', 'Rate limit (200 RPD on free tier): app suggests trying again later.'),
    ])

    # 2.5 Analysis Model
    heading2(doc, '2.5  Analysis Model')
    heading3(doc, '2.5.1  Analysis Class Diagram')
    body(doc,
         'The analysis class diagram captures the conceptual entities of the Mind Space '
         'system and their relationships, without implementation detail such as data types '
         'or method signatures. The central entities are: User, Chat, Message, Reflection, '
         'Arc, ArcInsight, and EmotionSpirit. Three service classes represent the AI '
         'subsystems: SafetyClassifier, EmbeddingService, and SageAI. The User entity '
         'has a one-to-many relationship with Chat and Arc. Each Chat has zero or one '
         'Reflection (created on Wrap Up). Each Reflection is associated with one Arc '
         'and one EmotionSpirit. Each Arc may have many Reflections and zero or one '
         'ArcInsight at any given time.')
    add_figure(doc, img('page25_img00.png'),
               'Figure 8 – Analysis Class Diagram', width_cm=13)

    # 2.6 Design Model
    heading2(doc, '2.6  Design Model')
    heading3(doc, '2.6.1  Design Class Diagram')
    body(doc,
         'The design class diagram refines the analysis model by introducing typed attributes, '
         'visibility modifiers, and key method signatures, reflecting the actual database '
         'schema and Flutter model classes. The most technically significant addition is the '
         'vector(1536) type on the reflections.embedding and arcs.centroid_embedding columns, '
         'which represents the 1536-dimensional float array produced by the OpenAI '
         'text-embedding-3-small model. The Arc entity includes processing_stage (forming | '
         'venting | stabilising | processing | shifting | integrating) and status (active | '
         'archived). The ArcInsight entity includes session_count_at_generation for cache '
         'invalidation logic.')
    add_figure(doc, img('page25_img00.png'),
               'Figure 9 – Design Class Diagram', width_cm=13)

    heading3(doc, '2.6.2  Sequence Diagrams')
    body(doc,
         'Five sequence diagrams describe the dynamic behaviour of the most critical system '
         'flows. Each is accompanied by a textual description table.')

    heading4(doc, '2.6.2.1  SD1 — User Authentication')
    add_figure(doc, img('page26_img00.png'),
               'Figure 10 – Sequence Diagram SD1: User Authentication', width_cm=13)
    body(doc, 'Table 11 – SD1: Textual Description — User Authentication', bold=True, indent=False)
    sd_table(doc, [
        ('Title', 'User Authentication'),
        ('Actors', 'Flutter App, Supabase Auth Service, PostgreSQL (profiles trigger)'),
        ('Precondition', 'User opens application for the first time or returns after sign-out'),
        ('Nominal Flow',
         '1. Flutter displays onboarding (first launch) → auth screen.\n'
         '2. User selects Google OAuth or magic link.\n'
         '3. Supabase Auth validates credentials → issues JWT.\n'
         '4. handle_new_user() trigger inserts profile row.\n'
         '5. Flutter stores JWT in secure storage → navigates to Home.'),
        ('Alternative', 'Returning user: JWT valid → direct to Home. Profile exists → trigger skipped.'),
        ('Postcondition', 'User authenticated; valid JWT in secure storage; profile row exists'),
    ])

    heading4(doc, '2.6.2.2  SD2 — Send Message with Safety Check')
    add_figure(doc, img('page28_img00.jpeg'),
               'Figure 11 – Sequence Diagram SD2: Send Message with Safety Check', width_cm=13)
    body(doc, 'Table 12 – SD2: Textual Description — Safety Check', bold=True, indent=False)
    sd_table(doc, [
        ('Title', 'Send Message with Safety Check'),
        ('Actors', 'Flutter App, safety-check Edge Function, Groq API (llama-3.1-8b-instant)'),
        ('Precondition', 'Active chat session; user has typed a message'),
        ('Nominal Flow',
         '1. User sends message.\n'
         '2. Flutter calls safety-check Edge Function (POST).\n'
         '3. Edge function calls Groq llama-3.1-8b-instant with temperature 0.1.\n'
         '4. Groq returns SAFE.\n'
         '5. Flutter proceeds to SD3 (chat-stream).'),
        ('Alternative (CRISIS)', 'Groq returns CRISIS → Flutter shows crisis resources card → message saved with flagged_for_safety = true → LLM not called.'),
        ('Error', 'API unavailable: system defaults to CRISIS (fail-closed design).'),
        ('Postcondition', 'Message classified; if SAFE → chat-stream initiated; if CRISIS → resources displayed'),
    ])

    heading4(doc, '2.6.2.3  SD3 — Chat Streaming via SSE')
    add_figure(doc, img('page30_img00.png'),
               'Figure 12 – Sequence Diagram SD3: Chat Streaming via SSE', width_cm=13)
    body(doc, 'Table 13 – SD3: Textual Description — Chat Streaming', bold=True, indent=False)
    sd_table(doc, [
        ('Title', 'Chat Streaming via Server-Sent Events (SSE)'),
        ('Actors', 'Flutter App, chat-stream Edge Function, Groq API (llama-3.3-70b-versatile), PostgreSQL'),
        ('Precondition', 'Message passed safety check (SAFE)'),
        ('Nominal Flow',
         '1. Flutter saves user message → calls chat-stream (POST with open HTTP connection).\n'
         '2. Edge function fetches last 20 messages from DB.\n'
         '3. Edge function opens streaming request to Groq (llama-3.3-70b, temp 0.7).\n'
         '4. Groq streams tokens as SSE events.\n'
         '5. Edge function forwards each token as SSE to Flutter.\n'
         '6. Flutter appends tokens to chat UI in real time.\n'
         '7. Stream completes → full assistant message saved to DB.'),
        ('Error', 'Groq unavailable: 503 returned → Flutter shows offline banner. Stream drop: partial buffer discarded.'),
        ('Postcondition', 'Full Sage reply displayed; both messages persisted in DB'),
    ])

    heading4(doc, '2.6.2.4  SD4 — Wrap Up: Reflection and Arc Assignment')
    add_figure(doc, img('page32_img00.png'),
               'Figure 13 – Sequence Diagram SD4: Wrap Up — Reflection Generation and Arc Assignment', width_cm=14)
    body(doc, 'Table 14 – SD4: Textual Description — Wrap Up', bold=True, indent=False)
    sd_table(doc, [
        ('Title', 'Wrap Up Session — Reflection Generation and Arc Assignment'),
        ('Actors', 'Flutter App, end-chat Edge Function, assign-arc Edge Function, Groq API, OpenAI Embeddings API, PostgreSQL'),
        ('Precondition', 'User taps Wrap Up; active chat has >= 1 exchange'),
        ('Nominal Flow',
         '1. Flutter calls end-chat (POST).\n'
         '2. end-chat fetches conversation + cross-arc context from DB.\n'
         '3. Groq llama-3.3-70b generates reflection JSON {spirit_id, what_sage_heard, question_to_sit_with, shared_perspective}.\n'
         '4. Reflection saved to reflections table.\n'
         '5. OpenAI text-embedding-3-small encodes what_sage_heard → vector(1536).\n'
         '6. end-chat calls assign-arc with the embedding.\n'
         '7. assign-arc computes cosine similarity (pgvector <=>) vs. all arc centroids.\n'
         '8a. similarity >= 0.78: assign to closest arc; update centroid (running average); increment session_count.\n'
         '8b. similarity < 0.65: create new arc; name with llama-3.1-8b-instant; set centroid = embedding.\n'
         '9. end-chat returns reflection_id + arc_id to Flutter.\n'
         '10. Flutter displays reflection card (cascade animation).'),
        ('Error', 'Reflection generation fails: retry shown. Embedding fails: reflection saved with needs_arc_assignment = true.'),
        ('Postcondition', 'Reflection and arc assignment persisted; reflection card displayed; chat status = completed'),
    ])

    heading4(doc, '2.6.2.5  SD5 — Generate Arc Insight')
    add_figure(doc, img('page34_img00.png'),
               'Figure 14 – Sequence Diagram SD5: Generate Arc Insight', width_cm=13)
    body(doc, 'Table 15 – SD5: Textual Description — Generate Arc Insight', bold=True, indent=False)
    sd_table(doc, [
        ('Title', 'Generate Arc Insight'),
        ('Actors', 'Flutter App, generate-arc-insight Edge Function, OpenRouter API (DeepSeek R1), PostgreSQL'),
        ('Precondition', 'User navigates to Arc Detail; Arc has >= 3 sessions; user taps Generate'),
        ('Nominal Flow',
         '1. Flutter calls generate-arc-insight (POST with arc_id).\n'
         '2. Edge function checks cache: valid insight exists AND session_count unchanged AND generated < 7 days ago → return cached.\n'
         '3. No valid cache: fetch all reflections in Arc (chronological order).\n'
         '4. Include cross-arc context (other active arcs).\n'
         '5. OpenRouter DeepSeek R1 (temp 0.3, max 1500 tokens) generates {how_it_evolved, pattern_noticed}.\n'
         '6. Insight saved with session_count_at_generation snapshot.\n'
         '7. Flutter displays Arc Insight screen.'),
        ('Alternative', 'Cache valid: return immediately without API call.'),
        ('Error', 'OpenRouter unavailable: retry button shown. Rate limit: "try again later" message.'),
        ('Postcondition', 'Insight displayed; result cached in arc_insights table'),
    ])

    # 2.7 Activity Model
    heading2(doc, '2.7  Activity Model')
    heading3(doc, '2.7.1  Arc Assignment Algorithm')
    body(doc,
         'The activity diagram below describes the complete Arc assignment algorithm, which '
         'constitutes the core technical contribution of this work. The algorithm runs as part '
         'of the assign-arc Supabase Edge Function, invoked by the end-chat function after '
         'the reflection embedding has been computed. The decision thresholds (0.78 and 0.65) '
         'were calibrated empirically during development. The centroid update formula is an '
         'incremental running average: '
         'new_centroid = (old_centroid × old_count + new_embedding) / (old_count + 1), '
         'which avoids recomputing the centroid over all historical embeddings on each update.')
    add_figure(doc, diag('fig_activity_diagram.png'),
               'Figure 15 – Activity Diagram: Arc Assignment Algorithm', width_cm=10)

    # 2.8 Implementation Model
    heading2(doc, '2.8  Implementation Model')
    heading3(doc, '2.8.1  Component Diagram')
    body(doc,
         'The component diagram presents the software architecture of Mind Space as a set '
         'of interconnected components distributed across three deployment nodes. On the '
         'Android device, the Flutter application is divided into three layers: the UI Layer '
         '(screens and widgets), the Riverpod State Layer (providers managing application '
         'state), and the Supabase Flutter Client (handling authentication, REST calls, and '
         'SSE streaming). On the Supabase Cloud node, the Auth Service manages JWT-based '
         'authentication, five Edge Functions implement the AI processing pipelines, and '
         'PostgreSQL with the pgvector extension provides the persistence layer. On the '
         'External AI Services node, three AI provider components are accessed via HTTPS: '
         'Groq (chat and safety classification), OpenAI (embeddings), and OpenRouter '
         '(arc insight generation).')
    add_figure(doc, diag('fig_component_diagram.png'),
               'Figure 16 – Component Diagram', width_cm=15)

    # 2.9 Deployment Model
    heading2(doc, '2.9  Deployment Model')
    heading3(doc, '2.9.1  Deployment Diagram')
    body(doc,
         'The deployment diagram maps the software artefacts onto their physical execution '
         'nodes. The Android Device node hosts the compiled Flutter application bundle (APK). '
         'The Supabase Cloud node hosts the PostgreSQL database (with pgvector), the Deno '
         'runtime for Edge Functions, the Auth server, and the PostgREST API layer. All five '
         'Edge Functions are deployed as serverless functions on this node. The three External '
         'AI Services nodes represent the Groq, OpenAI, and OpenRouter cloud platforms. '
         'All communication between the Android device and Supabase Cloud uses HTTPS '
         '(TLS 1.3). All communication between Supabase Edge Functions and the external AI '
         'services also uses HTTPS. No direct communication occurs between the Android '
         'device and the external AI service providers.')
    add_figure(doc, img('page36_img00.png'),
               'Figure 17 – Deployment Diagram', width_cm=14)

    # 2.10 Conclusion
    heading2(doc, '2.10  Conclusion')
    body(doc,
         'This chapter has produced a complete six-model RUP specification of the Mind Space '
         'system. The Use Case model defines fifteen use cases across three actor packages, '
         'each described in a standardised textual format. The Analysis and Design models '
         'capture the data structure and class relationships in increasing detail. The five '
         'Sequence Diagrams describe the most critical dynamic interactions, including the '
         'fail-closed safety mechanism and the token-by-token streaming flow. The Activity '
         'Diagram formalises the Arc assignment algorithm at the heart of the application\'s '
         'technical contribution. The Component and Deployment Diagrams complete the picture '
         'by showing the physical and logical organisation of the system\'s software '
         'components. Chapter 3 will document the implementation of this architecture.')
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # CHAPTER 3
    # ══════════════════════════════════════════════════════════
    heading1(doc, 'Chapter 3 – Implementation')

    # 3.1 Introduction
    heading2(doc, '3.1  Introduction')
    body(doc,
         'This chapter documents the implementation of Mind Space. We first present the '
         'development environment and the key technologies and frameworks employed. We '
         'then provide a comprehensive screen-by-screen presentation of the completed '
         'application, demonstrating that all requirements specified in Chapter 1 have '
         'been implemented.')

    # 3.2 Technologies
    heading2(doc, '3.2  Development Environment and Technologies')

    heading3(doc, '3.2.1  Flutter and Dart')
    body(doc,
         'Flutter is an open-source UI framework developed by Google that enables the '
         'construction of natively compiled applications for Android, iOS, Web, and desktop '
         'from a single Dart codebase. Mind Space targets Android for its initial version. '
         'Flutter\'s widget-based architecture and its hot-reload development cycle made it '
         'particularly well-suited to a rapid-iteration design process. The application\'s '
         'reactive state is managed entirely through flutter_riverpod, a compile-safe '
         'provider-based state management solution. Navigation is handled by go_router, '
         'which provides declarative URL-based routing with support for deep links and '
         'full-screen modal transitions. The choice of Flutter over React Native was driven '
         'by superior animation performance and the availability of fl_chart, a native Dart '
         'charting library used for the emotion journey graph in the Arc Detail screen.')

    heading3(doc, '3.2.2  Supabase')
    body(doc,
         'Supabase is an open-source Backend-as-a-Service platform built on top of PostgreSQL. '
         'It provides authentication (supporting Google OAuth and email magic links out of the '
         'box), a PostgREST-based auto-generated REST API, real-time subscriptions, and a '
         'serverless Edge Functions runtime based on Deno (TypeScript). Mind Space uses '
         'Supabase as its exclusive backend platform. Row-Level Security (RLS) policies on '
         'all tables ensure that every database query is automatically scoped to the '
         'authenticated user\'s ID, eliminating entire categories of authorisation bugs. '
         'The five Edge Functions (chat-stream, safety-check, end-chat, assign-arc, '
         'generate-arc-insight) are written in TypeScript and deployed on the Supabase '
         'edge runtime.')

    heading3(doc, '3.2.3  PostgreSQL and pgvector')
    body(doc,
         'PostgreSQL, accessed via Supabase, provides the relational persistence layer. '
         'The pgvector extension adds native support for storing and querying '
         'high-dimensional vectors, enabling the Arc assignment algorithm to compute '
         'cosine similarity directly in SQL using the <=> operator. An IVFFlat index '
         '(lists=100) on the reflections.embedding column ensures that approximate '
         'nearest-neighbour search scales efficiently as the number of reflections grows. '
         'The use of pgvector eliminates the need for a separate vector database service '
         '(such as Pinecone or Weaviate), keeping the architecture simple and fully '
         'within the Supabase platform.')

    heading3(doc, '3.2.4  Groq API')
    body(doc,
         'Groq provides inference-as-a-service for open-weight LLMs, optimised for speed '
         'using their proprietary Language Processing Unit (LPU) hardware. Mind Space uses '
         'two Groq-hosted models: llama-3.3-70b-versatile for conversation and reflection '
         'generation (temperature 0.7, max_tokens 1000), and llama-3.1-8b-instant for '
         'safety classification and arc naming (temperature 0.1 for classification, '
         '0.3 for naming). Groq\'s free tier provides over 1,000 requests per day on '
         'the 70B model and 14,400 requests per day on the 8B model, which is sufficient '
         'for the development and evaluation scope of this project. The OpenAI-compatible '
         'API interface simplifies integration with the edge functions.')

    heading3(doc, '3.2.5  OpenAI Embeddings API')
    body(doc,
         'The OpenAI text-embedding-3-small model is used exclusively for computing '
         'semantic embeddings of post-session reflection summaries. The model produces '
         '1536-dimensional floating-point vectors and is accessed via the standard '
         'OpenAI API endpoint. At the time of implementation, the cost is $0.02 per '
         'million input tokens, making it economically viable even at scale. The choice '
         'of text-embedding-3-small over the larger text-embedding-3-large was based on '
         'benchmark results showing minimal quality degradation on short emotional text '
         'classification tasks, at significantly lower cost and latency.')

    heading3(doc, '3.2.6  OpenRouter and DeepSeek R1')
    body(doc,
         'OpenRouter is an API gateway that provides unified access to a wide range of '
         'language models from different providers. Mind Space uses the '
         'deepseek/deepseek-r1:free model via OpenRouter exclusively for arc insight '
         'generation. DeepSeek R1 is a reasoning-optimised model that demonstrates '
         'strong performance on long-context synthesis tasks, making it well-suited to '
         'the arc insight use case, which requires reading all sessions in an Arc and '
         'producing a coherent longitudinal analysis. The free tier provides 200 requests '
         'per day, and the result caching strategy (7-day cache per Arc) ensures that '
         'this limit is rarely approached in practice.')

    heading3(doc, '3.2.7  Figma and Visual Studio Code')
    body(doc,
         'The visual design of Mind Space was produced in Figma, a collaborative browser-based '
         'interface design tool. All screens were designed to pixel-perfect specification '
         'before implementation, using a custom design system with consistent colour tokens, '
         'spacing, and typography (Inter font via google_fonts). The implementation was '
         'carried out in Visual Studio Code with the Flutter and Dart extensions, Dart '
         'DevTools for performance profiling, and the Supabase CLI for local Edge Function '
         'development and testing.')

    # 3.3 Application Presentation
    heading2(doc, '3.3  Application Presentation')
    body(doc,
         'The following section presents the completed Mind Space application screen by '
         'screen, following the user journey from first launch to arc insight generation. '
         'All screenshots were captured from the Android build on a physical device.')

    heading3(doc, '3.3.1  Onboarding and Authentication')
    body(doc,
         'The onboarding flow consists of three full-screen slides introducing Mind Space\'s '
         'core concepts: conversation with Sage, the reflection card, and the Arc system. '
         'The slides are shown once, on first launch only. After the third slide, the user '
         'is presented with the authentication screen, which offers two options: Sign in '
         'with Google (OAuth) and Continue with email (magic link). No password is ever '
         'requested or stored.')
    add_figure(doc, img('page38_img00.jpeg'),
               'Figure 18 – Onboarding screens (slide 1, slide 2, slide 3)', width_cm=13)

    heading3(doc, '3.3.2  Home Screen')
    body(doc,
         'The home screen is the central hub of the application. It displays a personalised '
         'greeting ("Morning, Seya.") with the current date, a 2×2 grid of the user\'s most '
         'recent open Arcs, and the persistent bottom navigation bar. The fourth slot in '
         'the grid is always the "New Thread" card, which opens a new free-format '
         'conversation. An ambient orb animation in the background reflects the user\'s '
         'overall emotional tone. The gradient background (deep purple to near-black) is '
         'applied consistently on every screen.')
    add_figure(doc, img('page39_img00.jpeg'),
               'Figure 19 – Home screen and Arc grid', width_cm=8)

    heading3(doc, '3.3.3  Chat with Sage — Free Mode')
    body(doc,
         'The chat screen slides up as a full-screen modal, covering the bottom navigation '
         'bar. In free mode (no Arc context), the session header reads "NEW SESSION · JUST '
         'TALKING" and the opening title is "What\'s with you, right now?" — an invitation '
         'to speak freely. Three quick-reply chips appear below Sage\'s opening message, '
         'offering: "Something on my mind", "Just want to vent", and "I don\'t really know". '
         'These chips disappear permanently after the first message is sent. Sage\'s messages '
         'appear on the left with a spirit orb beside them; user messages appear on the '
         'right as solid purple bubbles. The input bar is fixed at the bottom of the screen '
         'and rises with the keyboard.')
    add_figure(doc, img('page40_img00.jpeg'),
               'Figure 20 – Chat screen in free mode (with quick-reply chips)', width_cm=7)
    add_figure(doc, img('page40_img01.jpeg'),
               'Figure 21 – Chat screen in free mode (ongoing conversation)', width_cm=7)

    heading3(doc, '3.3.4  Chat with Sage — Arc Mode')
    body(doc,
         'When a chat is opened from an existing Arc card, it enters Arc mode. The header '
         'pill displays the Arc folder icon and name instead of the dashed "no arc yet" '
         'indicator. The session counter and Arc name appear in the session header ('
         '"SESSION 6 · THE JOB HUNT"). A "Last time" card appears above the first '
         'message, showing the what_sage_heard summary from the most recent session '
         'in this Arc — providing immediate context continuity. Sage\'s prompting strategy '
         'is modulated by the Arc\'s current EFT processing stage.')
    add_figure(doc, img('page40_img02.jpeg'),
               'Figure 22 – Chat screen in Arc mode with carry-forward card', width_cm=7)

    heading3(doc, '3.3.5  Post-Session Reflection Card')
    body(doc,
         'The reflection card is the emotional core of the application — the moment that '
         'transforms a conversation into insight. After Wrap Up, the card appears with a '
         'cascade entrance animation. It displays: the EFT spirit archetype (a glowing '
         'animated orb with its name — e.g., "Anxious"), the what_sage_heard text (one '
         'sentence capturing the primary emotional truth, not the surface label), the '
         'question_to_sit_with (an open, non-prescriptive question), and the '
         'shared_perspective (a normalising "many people find..." framing). '
         'The reflection is stored permanently and is accessible from the History Timeline '
         'and Arc Detail screens.')
    add_figure(doc, img('page41_img00.jpeg'),
               'Figure 23 – Post-session reflection card', width_cm=8)

    heading3(doc, '3.3.6  History — Timeline and Arcs Views')
    body(doc,
         'The History screen contains two tabs: Timeline and Arcs. The Timeline tab '
         'presents all past sessions in reverse-chronological order, grouped by date. '
         'Each session card shows the Arc pill, the session time and duration, a '
         '"View Reflection →" button, and the first line of the what_sage_heard text '
         'in italics. Filter chips above the list allow filtering by emotion spirit. '
         'The Arcs tab (also accessible from the home screen\'s "view all ›" link) '
         'shows the full grid of active Arcs and a collapsible section for archived Arcs, '
         'which are displayed at 50% opacity with a folder_archived.png icon.')
    add_figure(doc, img('page40_img03.jpeg'),
               'Figure 24 – History Timeline view (left) and Arcs grid view (right)', width_cm=12)

    heading3(doc, '3.3.7  Arc Detail and Arc Insight')
    body(doc,
         'The Arc Detail screen provides a complete view of a single Arc. The header '
         'shows the Arc name (tappable for inline rename), the ACTIVE/ARCHIVED status, '
         'and key statistics (number of sessions, start date, total span in days). '
         'An emotion journey chart (fl_chart LineChart) plots the dominant spirit of each '
         'session chronologically, with peak markers. Below the chart, each session is '
         'listed as a card with its title, date, a quote from the reflection, and an '
         'emotion tag. A sticky "Continue this Arc →" button anchors to the bottom '
         'of the screen.')
    add_figure(doc, img('page41_img01.jpeg'),
               'Figure 25 – Arc Detail screen with emotion journey graph', width_cm=8)
    body(doc,
         'The Arc Insight screen (accessible via the "Generate Arc Analysis" CTA, '
         'which appears when session_count ≥ 3) presents the longitudinal macro-analysis '
         'generated by DeepSeek R1. The screen displays how_it_evolved (a paragraph '
         'describing the emotional journey across all sessions) and pattern_noticed '
         '(a single observation about a recurring element). A dominant emotional weight '
         'bar shows the proportion of each spirit across all sessions in the Arc.')
    add_figure(doc, img('page42_img00.jpeg'),
               'Figure 26 – Arc Insight (Arc Analysis) screen', width_cm=8)

    heading3(doc, '3.3.8  Settings and Profile')
    body(doc,
         'The Settings screen provides the user with control over their account and '
         'preferences. A profile card at the top shows the user\'s display name, '
         'session count, number of active Arcs, and current reflection streak. '
         'The settings rows are organised into two sections: "YOUR SPACE" (containing '
         'Sage\'s tone, quiet hours, and daily check-in nudge — all marked "Coming soon" '
         'for v1) and "PRIVACY" (containing Sign Out, Export Data as JSON, and Delete '
         'Account — all functional in v1). The destructive actions (sign out, delete) '
         'include confirmation dialogs.')
    add_figure(doc, img('page42_img01.jpeg'),
               'Figure 27 – Profile and Settings screen', width_cm=8)

    # 3.4 Conclusion
    heading2(doc, '3.4  Conclusion')
    body(doc,
         'This chapter has presented the complete implementation of Mind Space. The '
         'technology stack — Flutter, Riverpod, GoRouter, Supabase Edge Functions, '
         'PostgreSQL with pgvector, Groq, OpenAI Embeddings, and OpenRouter — provides '
         'a modern, serverless, privacy-first architecture that satisfies all twelve '
         'functional requirements and all six non-functional requirements defined in '
         'Chapter 1. The screen-by-screen presentation demonstrates that the full user '
         'journey — from onboarding through conversation, reflection, arc evolution, '
         'and longitudinal insight — has been implemented and is functional on Android. '
         'The General Conclusion that follows will summarise the contributions of this '
         'work and identify directions for future development.')
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # GENERAL CONCLUSION
    # ══════════════════════════════════════════════════════════
    p = doc.add_paragraph(style='Titre 1 sans numéro')
    p.add_run('General Conclusion').font.name = 'Times New Roman'

    body(doc,
         'This project has designed, modelled, and implemented Mind Space, an AI-powered '
         'mobile application for guided emotional self-reflection on Android. Starting from '
         'the observation that the global mental health treatment gap is large, structural, '
         'and unlikely to be resolved by scaling clinical services alone, we proposed a '
         'technology-assisted approach grounded in established psychological frameworks — '
         'specifically Emotion-Focused Therapy and Pennebaker\'s expressive writing research '
         '— rather than in CBT protocol scripts.')
    body(doc,
         'The primary technical contribution of this work is the Arc system: an '
         'embedding-based automatic thematic clustering mechanism that groups a user\'s '
         'conversation sessions into persistent, evolving emotional themes without any '
         'manual categorisation. The system uses the OpenAI text-embedding-3-small model '
         'to encode post-session reflection summaries as 1536-dimensional semantic vectors, '
         'and computes cosine similarity against Arc centroids stored natively in PostgreSQL '
         'via the pgvector extension. A calibrated threshold of 0.78 determines assignment. '
         'The centroid is updated incrementally using a running average formula, keeping '
         'the algorithm efficient as Arc histories grow. Over time, this mechanism allows '
         'users to observe how their recurring emotional themes have evolved — a form of '
         'longitudinal self-knowledge that no existing consumer application provides.')
    body(doc,
         'The secondary contribution is the safety architecture: a pre-message binary '
         'classifier (SAFE/CRISIS) running on a separate, faster language model at '
         'near-zero temperature, designed to fail closed if the API is unavailable. '
         'This design eliminates the risk of the main conversational model responding '
         'to crisis content in an uncontrolled way, and provides contextually appropriate '
         'resources — including the Algerian emergency number SAMU 15 — when distress '
         'is detected. No prior consumer mental health application in the Algerian market '
         'has published a comparable safety architecture.')
    body(doc,
         'This work has its limits. Version 1.0 targets Android only; iOS support requires '
         'additional platform-specific authentication configuration. The Arc assignment '
         'thresholds (0.78 and 0.65) were calibrated empirically during development rather '
         'than validated on a large labelled dataset. The safety classifier has not been '
         'evaluated against a standardised clinical benchmark. The application has not '
         'undergone a formal user study; evidence of therapeutic efficacy is out of scope '
         'for a licence project and would require clinical collaboration. The free-tier '
         'API limits of Groq, OpenAI, and OpenRouter are adequate for development but '
         'would need to be revisited for a production deployment at scale.')
    body(doc,
         'Future work directions include: (1) the implementation of voice input and '
         'voice-to-text transcription using the mic button already present in the UI; '
         '(2) Arc merging logic for clusters whose centroids drift within 0.85 cosine '
         'similarity over time; (3) a cross-arc case formulation view mapping '
         'predisposing, precipitating, perpetuating, and protective factors; '
         '(4) a formal usability study and safety classifier evaluation; '
         '(5) an iOS build and App Store deployment. Mind Space stands as a working '
         'demonstration that a small engineering team can build a psychologically '
         'grounded, technically sophisticated, privacy-respecting mental health '
         'companion using exclusively free and open-source tools — and that technology, '
         'applied with care, can make a meaningful difference in the lives of people '
         'who feel unseen in their emotions.')
    add_page_break(doc)

    # ══════════════════════════════════════════════════════════
    # BIBLIOGRAPHY
    # ══════════════════════════════════════════════════════════
    p = doc.add_paragraph(style='Titre 1 sans numéro')
    p.add_run('Bibliography').font.name = 'Times New Roman'
    doc.add_paragraph()

    refs = [
        '[1]  World Health Organisation, World Mental Health Report: Transforming Mental Health for All. Geneva: WHO, 2022.',
        '[2]  L. S. Greenberg, "Emotion-focused therapy," Clinical Psychology and Psychotherapy, vol. 11, no. 1, pp. 3–16, 2004.',
        '[3]  J. W. Pennebaker, "Writing about emotional experiences as a therapeutic process," Psychological Science, vol. 8, no. 3, pp. 162–166, 1997.',
        '[4]  P. Ekman, "An argument for basic emotions," Cognition and Emotion, vol. 6, no. 3–4, pp. 169–200, 1992.',
        '[5]  J. A. Russell, "A circumplex model of affect," Journal of Personality and Social Psychology, vol. 39, no. 6, pp. 1161–1178, 1980.',
        '[6]  P. Kruchten, The Rational Unified Process: An Introduction, 3rd ed. Boston: Addison-Wesley, 2003.',
        '[7]  G. Booch, J. Rumbaugh, and I. Jacobson, The Unified Modeling Language User Guide, 2nd ed. Boston: Addison-Wesley, 2005.',
        '[8]  IEEE Std 830-1998, IEEE Recommended Practice for Software Requirements Specifications. New York: IEEE, 1998.',
        '[9]  A. Inkster, S. Sarda, and V. Subramanian, "An empathy-driven, conversational AI agent for digital mental health: a mixed methods pilot study," JMIR Mental Health, vol. 5, no. 4, p. e12106, 2018.',
        '[10] K. K. Fitzpatrick, A. Darcy, and M. Vierhile, "Delivering cognitive behavior therapy to young adults with symptoms of depression and anxiety using a fully automated conversational agent (Woebot): a randomized controlled trial," JMIR Mental Health, vol. 4, no. 2, p. e19, 2017.',
        '[11] Meta AI, "Llama 3.3 70B Versatile," Groq inference API, 2024. [Online]. Available: https://groq.com/',
        '[12] OpenAI, "Text Embedding Models," OpenAI API documentation, 2024. [Online]. Available: https://platform.openai.com/docs/guides/embeddings',
        '[13] DeepSeek AI, "DeepSeek R1," OpenRouter, 2025. [Online]. Available: https://openrouter.ai/',
        '[14] Google LLC, "Flutter — Build apps for any screen," Flutter framework, 2024. [Online]. Available: https://flutter.dev/',
        '[15] Supabase Inc., "Supabase — The Open Source Firebase Alternative," 2024. [Online]. Available: https://supabase.com/',
        '[16] A. Timkovskiy, "pgvector: Open-source vector similarity search for Postgres," GitHub, 2024. [Online]. Available: https://github.com/pgvector/pgvector',
    ]
    for ref in refs:
        p = doc.add_paragraph(style='bibliographie')
        run = p.add_run(ref)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        p.paragraph_format.left_indent = Cm(1.0)

    # ── Save ──────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f'\nReport saved to: {OUT_PATH}')
    print(f'  Size: {os.path.getsize(OUT_PATH) / 1024:.0f} KB')


if __name__ == '__main__':
    build()
