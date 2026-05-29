from docx import Document
from docx.oxml.ns import qn

doc = Document(r'd:\program\mind_space\cahierDecharge\MindSpace_PFE_Report.docx')

print('=== HEADING STRUCTURE ===')
for p in doc.paragraphs:
    style = p.style.name
    if ('Heading' in style or 'Titre' in style) and p.text.strip():
        print(f'  [{style}] {p.text[:80]}')

print()
print(f'Total paragraphs: {len(doc.paragraphs)}')
print(f'Total tables:     {len(doc.tables)}')

imgs = doc.element.body.findall('.//' + qn('w:drawing'))
print(f'Total images:     {len(imgs)}')

print()
print('=== FIRST 5 TABLE STRUCTURES ===')
for i, t in enumerate(doc.tables[:5]):
    print(f'  Table {i+1}: {len(t.rows)} rows x {len(t.columns)} cols')
    if t.rows:
        header = [c.text[:30] for c in t.rows[0].cells]
        print(f'    Header: {header}')
